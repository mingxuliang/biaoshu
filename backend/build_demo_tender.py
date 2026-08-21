"""生成一份用于 P1（招标文件解析与评标尺子锁定）端到端测试的示例招标文件。

包含明确的评分规则、必响应/否决条款、资格要求、格式要求，以及可被 vetoParams
抽取的关键数值（投标有效期、预算上限、资产负债率上限），便于验证 E0 解析引擎
与 E1/E2 尺子驱动逻辑是否生效。
"""

import os

from docx import Document

OUT_PATH = os.path.join(os.path.dirname(__file__), "sample_data", "demo_tender.docx")


def build() -> None:
    doc = Document()

    doc.add_heading("企业培训管理系统建设项目招标文件", level=1)

    doc.add_heading("第一章 投标人须知", level=2)
    doc.add_paragraph("1.1 项目预算：本项目预算上限为人民币 500 万元（含税），投标报价超过预算上限的按无效标处理。")
    doc.add_paragraph("1.2 投标有效期：投标有效期自投标截止之日起不少于 90 日历天，逾期视为无效。")
    doc.add_paragraph(
        "1.3 财务要求：投标人须提交近三年经审计财务报表，资产负债率不高于 85%，超过该比例须附专项说明函。"
    )
    doc.add_paragraph(
        "1.4 资格要求：投标人须提供有效的营业执照、软件企业认定证书、ISO 9001 质量管理体系认证证书；"
        "拟派项目经理须具备 PMP 或信息系统项目管理师（高级）资格。"
    )
    doc.add_paragraph("1.5 暗标评审：本项目为暗标评审，技术标不得出现任何可识别投标人身份的标记、名称或LOGO。")

    doc.add_heading("第二章 评标办法及评分标准", level=2)
    doc.add_paragraph("综合评分法，满分 100 分，评分维度及权重如下：")
    doc.add_paragraph("2.1 技术方案（35 分）：系统架构合理性、功能响应完整性、技术先进性。")
    doc.add_paragraph("2.2 项目团队（15 分）：项目经理资质、核心人员经验、团队稳定性。")
    doc.add_paragraph("2.3 商务报价（25 分）：报价合理性、成本控制措施、付款方式接受度。")
    doc.add_paragraph("2.4 实施与售后（15 分）：实施计划完整性、培训方案、售后服务承诺。")
    doc.add_paragraph("2.5 演示答辩（10 分）：现场演示表现、问题响应能力。")

    doc.add_heading("第三章 废标与无效投标情形", level=2)
    doc.add_paragraph("3.1 投标报价超过预算上限的，按废标处理。")
    doc.add_paragraph("3.2 投标有效期不满足 90 日历天要求的，按废标处理。")
    doc.add_paragraph("3.3 技术标出现可识别投标人身份标记的（暗标评审要求），按废标处理。")
    doc.add_paragraph("3.4 未按要求提交全部资格证明文件的，按无效投标处理。")

    doc.add_heading("第四章 响应文件格式要求", level=2)
    doc.add_paragraph("4.1 投标文件正本 1 份、副本 5 份，采用 A4 纸左侧装订，统一封面模板。")
    doc.add_paragraph("4.2 正文宋体小四、行距 1.5 倍，全文连续页码，目录须与正文页码一致。")
    doc.add_paragraph("4.3 法定代表人或授权代表须在指定位置签字并加盖单位公章，骑缝章须覆盖全部页边。")

    os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
    doc.save(OUT_PATH)
    print(f"已生成：{OUT_PATH}")


if __name__ == "__main__":
    build()
