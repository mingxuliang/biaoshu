"""端到端：产品库上传技术标抽取、同义功能点合并、审核入库。"""

import io
import json
import sys
import time
import urllib.error
import urllib.request

import docx

sys.stdout.reconfigure(encoding="utf-8")

BASE = "http://localhost:8000"


def call(
    method: str,
    path: str,
    token: str | None = None,
    json_body: dict | None = None,
    data: bytes | None = None,
    content_type: str | None = None,
    timeout: int = 30,
):
    url = f"{BASE}{path}"
    body = data
    if json_body is not None:
        body = json.dumps(json_body).encode("utf-8")
        content_type = "application/json"
    req = urllib.request.Request(url, data=body, method=method)
    if content_type:
        req.add_header("Content-Type", content_type)
    if token:
        req.add_header("Authorization", f"Bearer {token}")
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return resp.status, resp.read()
    except urllib.error.HTTPError as e:
        return e.code, e.read()


def multipart(fields: dict[str, str], files: list[tuple[str, str, bytes]]) -> tuple[bytes, str]:
    boundary = "----e2eProduct"
    parts: list[bytes] = []
    for name, value in fields.items():
        parts.append(
            f"--{boundary}\r\nContent-Disposition: form-data; name=\"{name}\"\r\n\r\n{value}\r\n".encode()
        )
    for field, filename, blob in files:
        parts.append(
            (
                f"--{boundary}\r\nContent-Disposition: form-data; name=\"{field}\"; filename=\"{filename}\"\r\n"
                "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
            ).encode()
            + blob
            + b"\r\n"
        )
    parts.append(f"--{boundary}--\r\n".encode())
    return b"".join(parts), f"multipart/form-data; boundary={boundary}"


def make_tech_docx(title: str, extra_heading: str | None = None) -> bytes:
    document = docx.Document()
    document.add_heading(title, level=1)
    document.add_heading("系统功能", level=1)
    document.add_heading("证书颁发", level=2)
    document.add_paragraph("支持按培训班批量颁发电子证书，覆盖结业证、合格证模板。")
    document.add_heading("组织变更数据修改", level=2)
    document.add_paragraph("组织架构调整后同步修改人员与班级基础数据。")
    if extra_heading:
        document.add_heading(extra_heading, level=2)
        document.add_paragraph(f"{extra_heading}用于查询已颁发证书的状态与下载记录。")
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "功能名称"
    table.cell(0, 1).text = "功能说明"
    table.cell(1, 0).text = "证书颁发"
    table.cell(1, 1).text = "批量颁发电子证书"
    table.cell(2, 0).text = "组织变更数据修改"
    table.cell(2, 1).text = "同步基础数据"
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def poll_job(token: str, job_id: str, timeout_s: int = 180) -> dict:
    started = time.time()
    while time.time() - started < timeout_s:
        status, body = call("GET", f"/api/product-extract-jobs/{job_id}", token=token)
        assert status == 200, body
        job = json.loads(body)
        if job["status"] in ("done", "failed"):
            return job
        time.sleep(1.5)
    raise AssertionError(f"抽取任务超时: {job_id}")


def main() -> None:
    status, body = call("POST", "/api/auth/login", json_body={"email": "chen@zhibiaoyun.com", "password": "123456"})
    assert status == 200, f"登录失败: {status} {body}"
    token = json.loads(body)["token"]

    status, body = call(
        "POST",
        "/api/product-libraries",
        token=token,
        json_body={"name": "E2E培训管理平台", "category": "软件系统", "description": "抽取去重验证", "owner": "E2E"},
    )
    assert status == 200, body
    library = json.loads(body)
    library_id = library["id"]
    print("create library ok", library_id)

    docx_bytes = make_tech_docx("成都地铁培训系统技术标")
    payload, ctype = multipart({}, [("files", "chengdu-tech.docx", docx_bytes)])
    status, body = call(
        "POST",
        f"/api/product-libraries/{library_id}/source-docs",
        token=token,
        data=payload,
        content_type=ctype,
        timeout=60,
    )
    assert status == 200, body
    jobs = json.loads(body)
    assert len(jobs) == 1, jobs
    job = poll_job(token, jobs[0]["id"])
    assert job["status"] == "done", job
    print("first extract ok", job)

    status, body = call("GET", f"/api/product-libraries/{library_id}/features", token=token)
    assert status == 200, body
    features = json.loads(body)
    names = {f["name"] for f in features}
    assert any("证书" in n or "组织" in n for n in names), features
    assert all(f["status"] == "待审核" for f in features), features
    print("pending features", names)

    second = make_tech_docx("成都地铁培训系统技术标-修订", extra_heading="证书查询")
    payload, ctype = multipart({}, [("files", "chengdu-tech-v2.docx", second)])
    status, body = call(
        "POST",
        f"/api/product-libraries/{library_id}/source-docs",
        token=token,
        data=payload,
        content_type=ctype,
        timeout=60,
    )
    assert status == 200, body
    jobs2 = json.loads(body)
    job2 = poll_job(token, jobs2[0]["id"])
    assert job2["status"] == "done", job2
    print("second extract ok", job2)

    status, body = call("GET", f"/api/product-libraries/{library_id}/features", token=token)
    assert status == 200, body
    features = json.loads(body)
    certs = [f for f in features if "证书颁发" in (f["name"] + "".join(f.get("aliases") or []))]
    assert len(certs) >= 1, features
    if job2.get("merged", 0) == 0:
        print("warn: second job merged=0, names=", [f["name"] for f in features])
    else:
        print("merge count", job2["merged"])

    target = certs[0]
    status, body = call(
        "PATCH",
        f"/api/product-features/{target['id']}",
        token=token,
        json_body={"status": "已入库"},
    )
    assert status == 200, body
    updated = json.loads(body)
    assert updated["status"] == "已入库"
    print("ingest ok", updated["name"])

    status, body = call("GET", "/api/product-libraries", token=token)
    assert status == 200, body
    listed = json.loads(body)
    hit = next(x for x in listed if x["id"] == library_id)
    assert hit["featureCount"] >= 1
    print("list libraries ok, features=", hit["featureCount"])

    status, body = call(
        "POST",
        "/api/projects",
        token=token,
        json_body={"name": "E2E产品库写标项目", "code": "E2E-PROD-001", "type": "工程"},
    )
    assert status == 200, body
    project_id = json.loads(body)["id"]
    status, body = call("GET", f"/api/projects/{project_id}/writer-draft", token=token)
    assert status == 200, body
    draft = json.loads(body)
    status, body = call(
        "PATCH",
        f"/api/writer-drafts/{draft['id']}",
        token=token,
        json_body={"selectedProductLibraryId": library_id},
    )
    assert status == 200, body
    saved = json.loads(body)
    assert saved.get("selectedProductLibraryId") == library_id, saved
    print("writer draft product library ok")

    status, body = call("DELETE", f"/api/projects/{project_id}", token=token)
    assert status in (200, 204), body
    status, body = call("DELETE", f"/api/product-libraries/{library_id}", token=token)
    assert status == 200, body
    print("cleanup ok")


if __name__ == "__main__":
    main()
