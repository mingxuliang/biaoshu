"""端到端：招标原文段落、解析报告导出、项目进度由真实节点派生。"""

import io
import json
import sys
import urllib.error
import urllib.request

import docx

sys.stdout.reconfigure(encoding="utf-8")

BASE = "http://localhost:8000"


def call(
    method: str,
    path: str,
    token: str | None = None,
    json_body: dict | None = None,
    data: bytes | None = None,
    content_type: str | None = None,
):
    url = f"{BASE}{path}"
    body = data
    if json_body is not None:
        body = json.dumps(json_body).encode("utf-8")
        content_type = "application/json"
    req = urllib.request.Request(url, data=body, method=method)
    if content_type:
        req.add_header("Content-Type", content_type)
    if token:
        req.add_header("Authorization", f"Bearer {token}")
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            return resp.status, resp.read()
    except urllib.error.HTTPError as e:
        return e.code, e.read()


def main() -> None:
    status, body = call("POST", "/api/auth/login", json_body={"email": "chen@zhibiaoyun.com", "password": "123456"})
    assert status == 200, f"登录失败: {status} {body}"
    token = json.loads(body)["token"]

    status, body = call("GET", "/api/projects", token=token)
    assert status == 200, body
    projects = json.loads(body)
    for p in projects:
        assert 0 <= int(p["progress"]) <= 100
        assert float(p["score"]) >= 0
    print("list projects live progress ok, count=", len(projects))

    status, body = call(
        "POST",
        "/api/projects",
        token=token,
        json_body={"name": "解析原文验证项目", "code": "E2E-PARSE-001", "type": "工程"},
    )
    assert status == 200, body
    project = json.loads(body)
    assert project["progress"] == 0
    assert project["score"] == 0
    project_id = project["id"]

    document = docx.Document()
    document.add_heading("企业培训管理系统招标文件", level=1)
    document.add_paragraph("招标编号：E2E-PARSE-001")
    document.add_heading("评标办法", level=1)
    document.add_paragraph("技术标 60 分，商务标 40 分。暗标评审不得出现投标人名称。")
    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    boundary = "----e2eParse"
    parts = [
        f"--{boundary}\r\nContent-Disposition: form-data; name=\"project_id\"\r\n\r\n{project_id}\r\n".encode(),
        (
            f"--{boundary}\r\nContent-Disposition: form-data; name=\"file\"; filename=\"e2e-parse.docx\"\r\n"
            "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
        ).encode()
        + docx_bytes
        + b"\r\n",
        f"--{boundary}--\r\n".encode(),
    ]
    status, body = call(
        "POST",
        "/api/tender-documents",
        token=token,
        data=b"".join(parts),
        content_type=f"multipart/form-data; boundary={boundary}",
    )
    assert status == 200, f"上传失败: {status} {body}"
    tender = json.loads(body)

    status, body = call("GET", f"/api/tender-documents/{tender['id']}/paragraphs", token=token)
    assert status == 200, f"抽取段落失败: {status} {body}"
    paras = json.loads(body)
    texts = " ".join(p["text"] for p in paras)
    assert "评标办法" in texts or "招标文件" in texts
    print("tender paragraphs ok, count=", len(paras))

    status, body = call("GET", f"/api/projects/{project_id}", token=token)
    assert status == 200, body
    refreshed = json.loads(body)
    assert refreshed["progress"] >= 15, f"上传招标文件后进度应增加，实际 {refreshed['progress']}"
    print("progress after tender upload=", refreshed["progress"])

    status, body = call("GET", f"/api/projects/{project_id}/checklist/latest", token=token)
    if status == 200:
        checklist = json.loads(body)
        if checklist.get("status") == "done":
            status, report = call(
                "GET",
                f"/api/projects/{project_id}/checklist/{checklist['id']}/export",
                token=token,
            )
            assert status == 200, report
            assert report[:2] == b"PK", "导出的不是 docx"
            print("checklist export ok, bytes=", len(report))
        else:
            print("checklist not done, skip export")
    else:
        exported = False
        for p in projects:
            status, body = call("GET", f"/api/projects/{p['id']}/checklist/latest", token=token)
            if status != 200:
                continue
            checklist = json.loads(body)
            if checklist.get("status") != "done":
                continue
            status, report = call(
                "GET",
                f"/api/projects/{p['id']}/checklist/{checklist['id']}/export",
                token=token,
            )
            assert status == 200, report
            assert report[:2] == b"PK", "导出的不是 docx"
            print("checklist export ok from existing project, bytes=", len(report))
            exported = True
            break
        if not exported:
            print("no checklist yet, skip export")

    print("PASS")


if __name__ == "__main__":
    main()
