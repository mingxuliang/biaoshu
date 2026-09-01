"""修改闭环 / 约定对照 / 标题识别 的轻量单测，不依赖数据库。"""

from __future__ import annotations

import io
import tempfile
import unittest
from pathlib import Path

import docx

from app.engines.e3_semantic import CHUNK_CHARS, MAX_REVIEW_CHARS, _split_chunks
from app.engines.e_parse_match import _unanswered, run as parse_match_run
from app.engines.revision_build import _heading_level, writeback_docx


class HeadingLevelTests(unittest.TestCase):
    def test_plain_numbered_body_is_not_h1(self) -> None:
        para = {"text": "1. 投标人应具备独立法人资格", "bold": False, "fontSizePt": 12}
        self.assertIsNone(_heading_level(para, 12.0))

    def test_bold_short_numbered_is_h1(self) -> None:
        para = {"text": "1. 编制说明", "bold": True, "fontSizePt": 16}
        self.assertEqual(_heading_level(para, 12.0), 1)

    def test_chapter_heading(self) -> None:
        para = {"text": "第一章 施工组织设计", "bold": False, "fontSizePt": 12}
        self.assertEqual(_heading_level(para, 12.0), 1)


class ParseMatchTests(unittest.TestCase):
    def test_title_in_bid_is_covered(self) -> None:
        self.assertFalse(
            _unanswered(
                "投标人须提供有效营业执照副本扫描件",
                "详见附件：本公司营业执照及资质证书",
                "营业执照",
            )
        )

    def test_heading_covers_dimension(self) -> None:
        self.assertFalse(
            _unanswered(
                "施工进度计划应明确关键节点与工期保证措施",
                "正文未展开细节",
                "进度计划",
                "第一章 进度计划\n第二章 质量保证",
            )
        )

    def test_single_short_window_is_uncovered(self) -> None:
        self.assertTrue(
            _unanswered(
                "投标人必须提交完全无关的专项核验材料并加盖鲜章",
                "本公司依法设立，具有独立承担民事责任能力",
                "专项核验材料",
            )
        )

    def test_score_rule_title_hit_skips_finding(self) -> None:
        findings = parse_match_run(
            "技术标目录含进度计划与质量保证体系",
            score_rules=[{"dimension": "进度计划", "detail": "应说明关键线路"}],
            strategy_keys={"checklist_map"},
            headings=["进度计划"],
        )
        self.assertEqual(findings, [])


class ChunkTests(unittest.TestCase):
    def test_short_text_is_one_chunk(self) -> None:
        chunks = _split_chunks("短文")
        self.assertEqual(len(chunks), 1)
        self.assertEqual(chunks[0]["text"], "短文")

    def test_long_text_covers_all_characters(self) -> None:
        blob = ("章节\n" + "正文" * 5000) * 8
        chunks = _split_chunks(blob)
        self.assertGreater(len(chunks), 4)
        joined = "\n".join(c["text"] for c in chunks)
        self.assertGreaterEqual(len(joined.replace("\n", "")), len(blob.replace("\n", "")) - 50)

    def test_prefers_chapter_breaks(self) -> None:
        chapter = "第一章 施工组织\n" + ("措施。" * 4000) + "\n第二章 质量保证\n" + ("验收。" * 4000)
        chunks = _split_chunks(chapter)
        self.assertGreaterEqual(len(chunks), 2)
        headings = " ".join(c["heading"] for c in chunks)
        self.assertTrue("第一章" in headings or "第二章" in headings)

    def test_chunk_size_stays_bounded(self) -> None:
        blob = "正文" * 20000
        chunks = _split_chunks(blob)
        for chunk in chunks[:-1]:
            self.assertLessEqual(len(chunk["text"]), CHUNK_CHARS + 100)

    def test_review_cap_is_about_300k(self) -> None:
        blob = ("章节标题足够长但不是标题\n" + "正文内容" * 8000) * 12
        chunks = _split_chunks(blob)
        total = sum(len(c["text"]) for c in chunks)
        self.assertLessEqual(len(chunks), 19)
        self.assertLessEqual(total, MAX_REVIEW_CHARS + 500)
        self.assertGreater(total, MAX_REVIEW_CHARS * 0.8)


class WritebackTests(unittest.TestCase):
    def test_writeback_replaces_text_keeps_bold(self) -> None:
        document = docx.Document()
        para = document.add_paragraph()
        run = para.add_run("甲方承诺投标有效期90天")
        run.bold = True
        with tempfile.TemporaryDirectory() as tmp:
            path = str(Path(tmp) / "src.docx")
            document.save(path)
            out = writeback_docx(path, [{"type": "paragraph", "text": "甲方承诺投标有效期120天"}])
        loaded = docx.Document(io.BytesIO(out))
        para = next(p for p in loaded.paragraphs if "120" in (p.text or ""))
        self.assertIn("120", para.text)
        self.assertTrue(any(run.bold for run in para.runs))


class ReviewExportTests(unittest.TestCase):
    def test_report_docx_contains_score_and_issues(self) -> None:
        from app.engines.review_export import review_run_to_docx

        blob = review_run_to_docx(
            project_name="华北演示",
            project_code="HB-001",
            round_no=3,
            overall=82.5,
            light="橙",
            waste=1,
            risk=4,
            suggest=2,
            levels=[{"key": "L1", "name": "一票否决扫描", "score": 90, "issues": 1, "status": "风险", "desc": "否决项"}],
            dimensions=[{"name": "完整性", "weight": 20, "score": 80}],
            issues=[
                {
                    "severity": "废标",
                    "rule": "未盖公章",
                    "level": "L1",
                    "location": "封面",
                    "excerpt": "投标函未加盖公章",
                    "tenderQuote": "投标文件须加盖公章",
                    "suggestion": "在投标函落款处加盖公章",
                }
            ],
        )
        loaded = docx.Document(io.BytesIO(blob))
        text = "\n".join(p.text for p in loaded.paragraphs)
        self.assertIn("AI 智能预审报告", text)
        self.assertIn("82.5", text)
        self.assertIn("未盖公章", text)
        self.assertIn("投标函未加盖公章", text)


if __name__ == "__main__":
    unittest.main()
