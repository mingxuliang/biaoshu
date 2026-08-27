"""端到端验证：撰写工作台真实编辑与导出。

1. 登录拿 token
2. 取/建撰写草稿，写入一个含加粗/斜体/下划线/列表的目录节点
3. 保存章节内容（真实落库）
4. 调用导出接口，拿到真实 .docx，用 python-docx 解析校验标题/加粗/列表
5. 确认 bid_documents 表新增一条 source=writer 记录，且 /projects/{id}/bid-documents 能查到
"""

import io
import json
import sys
import urllib.error
import urllib.request

sys.stdout.reconfigure(encoding="utf-8")

BASE = "http://localhost:8000"


def call(method: str, path: str, token: str | None = None, json_body: dict | None = None):
    url = f"{BASE}{path}"
    data = json.dumps(json_body).encode("utf-8") if json_body is not None else None
    req = urllib.request.Request(url, data=data, method=method)
    if json_body is not None:
        req.add_header("Content-Type", "application/json")
    if token:
        req.add_header("Authorization", f"Bearer {token}")
    try:
        with urllib.request.urlopen(req) as resp:
            body = resp.read()
            return resp.status, body
    except urllib.error.HTTPError as e:
        return e.code, e.read()


def main() -> None:
    status, body = call("POST", "/api/auth/login", json_body={"email": "chen@zhibiaoyun.com", "password": "123456"})
    assert status == 200, f"登录失败: {status} {body}"
    token = json.loads(body)["token"]
    print("login ok, token prefix:", token[:12])

    status, body = call(
        "POST",
        "/api/projects",
        token=token,
        json_body={"name": "撰写导出验证项目", "code": "E2E-WRITER-001", "type": "工程"},
    )
    assert status == 200, f"新建项目失败: {status} {body}"
    PROJECT_ID = json.loads(body)["id"]

    status, body = call("GET", f"/api/projects/{PROJECT_ID}/writer-draft", token=token)
    assert status == 200, f"获取草稿失败: {status} {body}"
    draft = json.loads(body)
    draft_id = draft["id"]
    print("draft id:", draft_id)

    outline = [
        {
            "id": "ch-1",
            "num": "1",
            "title": "项目理解与实施方案",
            "parentId": None,
            "expanded": True,
            "weight": 10,
            "dimension": None,
            "idea": "",
            "aiIdea": "",
            "optimized": False,
            "status": "已完成",
            "words": 0,
            "aiRounds": 0,
        },
        {
            "id": "ch-1-1",
            "num": "1.1",
            "title": "技术路线",
            "parentId": "ch-1",
            "expanded": True,
            "weight": 5,
            "dimension": None,
            "idea": "",
            "aiIdea": "",
            "optimized": False,
            "status": "已完成",
            "words": 0,
            "aiRounds": 0,
        },
    ]
    status, body = call(
        "PATCH",
        f"/api/writer-drafts/{draft_id}",
        token=token,
        json_body={"outline": outline},
    )
    assert status == 200, f"写入目录失败: {status} {body}"
    print("outline saved")

    chapter_content = (
        "本项目采用**装配式施工**工艺，显著提升现场作业效率。\n"
        "整体方案强调*安全可控*与__绿色环保__并重。\n"
        "- 第一阶段：土建施工\n"
        "- 第二阶段：设备安装与调试\n"
        "1. 组建专项管理团队\n"
        "2. 编制专项施工方案\n"
    )
    status, body = call(
        "PATCH",
        f"/api/writer-drafts/{draft_id}/chapters/ch-1-1",
        token=token,
        json_body={"content": chapter_content},
    )
    assert status == 200, f"保存章节内容失败: {status} {body}"
    print("chapter content saved")

    status, body = call("GET", f"/api/writer-drafts/{draft_id}/export", token=token)
    assert status == 200, f"导出失败: {status} {body}"
    docx_bytes = body
    print("export ok, size:", len(docx_bytes))

    out_path = "e2e_writer_export.docx"
    with open(out_path, "wb") as f:
        f.write(docx_bytes)

    import docx

    document = docx.Document(io.BytesIO(docx_bytes))
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
    print("headings:", headings)
    assert any("技术路线" in h for h in headings), "未找到章节标题"

    bold_found = False
    italic_found = False
    underline_found = False
    bullet_found = False
    for p in document.paragraphs:
        for run in p.runs:
            if run.bold and "装配式施工" in run.text:
                bold_found = True
            if run.italic and "安全可控" in run.text:
                italic_found = True
            if run.underline and "绿色环保" in run.text:
                underline_found = True
        if p.style.name == "List Bullet" and "土建施工" in p.text:
            bullet_found = True

    print("bold_found:", bold_found, "italic_found:", italic_found, "underline_found:", underline_found, "bullet_found:", bullet_found)
    assert bold_found and italic_found and underline_found and bullet_found, "行内样式/列表未正确还原"

    status, body = call("GET", f"/api/projects/{PROJECT_ID}/bid-documents", token=token)
    assert status == 200, f"获取项目文件列表失败: {status} {body}"
    docs = json.loads(body)
    writer_docs = [d for d in docs if d["source"] == "writer"]
    print("writer bid-documents count:", len(writer_docs))
    assert len(writer_docs) >= 1, "未找到 source=writer 的 BidDocument 记录"
    writer_doc_id = writer_docs[0]["id"]
    print("latest writer doc:", writer_docs[0]["filename"])

    # 验证「从已有投标文件中选择」→ 真实触发预审，而不是固定示例文档
    status, body = call(
        "POST",
        f"/api/projects/{PROJECT_ID}/prereview-jobs",
        token=token,
        json_body={"bid_document_id": writer_doc_id, "scope": "full"},
    )
    assert status == 200, f"创建预审任务失败: {status} {body}"
    job = json.loads(body)
    job_id = job["job_id"]
    print("prereview job created:", job_id)

    import time

    for _ in range(60):
        status, body = call("GET", f"/api/prereview-jobs/{job_id}", token=token)
        assert status == 200, f"查询任务状态失败: {status} {body}"
        job_status = json.loads(body)
        if job_status["status"] in ("done", "failed"):
            break
        time.sleep(2)
    print("prereview job status:", job_status["status"])
    assert job_status["status"] == "done", f"预审任务未成功完成: {job_status}"

    status, body = call("GET", f"/api/projects/{PROJECT_ID}/review-runs/latest", token=token)
    assert status == 200, f"获取最新预审报告失败: {status} {body}"
    report = json.loads(body)
    print("review report score:", report.get("score"), "issues:", len(report.get("issues", [])))

    print("\nALL CHECKS PASSED")


if __name__ == "__main__":
    main()
