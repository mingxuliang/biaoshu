"""把已完成的 AI 预审运行导出为可读 Word 报告。"""

from __future__ import annotations

import io
from datetime import datetime

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


def review_run_to_docx(
    *,
    project_name: str,
    project_code: str,
    round_no: int,
    overall: float,
    light: str,
    waste: int,
    risk: int,
    suggest: int,
    levels: list[dict],
    dimensions: list[dict],
    issues: list[dict],
    finished_at: datetime | None = None,
) -> bytes:
    document = docx.Document()
    title = document.add_heading("AI 智能预审报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    when = finished_at.strftime("%Y-%m-%d %H:%M") if finished_at else "—"
    meta = document.add_paragraph()
    meta.add_run(f"项目：{project_name or '（未命名）'}\n")
    meta.add_run(f"招标编号：{project_code or '—'}\n")
    meta.add_run(f"预审轮次：第 {round_no} 轮\n")
    meta.add_run(f"生成时间：{when}\n")
    meta.add_run("预审方法：L1-L5 分层预审")

    document.add_heading("一、综合预审得分与分层指标", level=1)
    p = document.add_paragraph()
    p.add_run(f"综合预审得分：{overall} 分（满分 100）").bold = True
    document.add_paragraph(
        f"风险灯：{light or '—'}　　废标 {waste} 项　　扣分 {risk} 项　　建议 {suggest} 项"
    )

    for level in levels or []:
        name = level.get("name") or level.get("key") or "—"
        score = level.get("score", "—")
        issues_n = level.get("issues", 0)
        status = level.get("status") or ""
        desc = level.get("desc") or ""
        key = level.get("key") or ""
        line = document.add_paragraph()
        line.add_run(f"{key} {name}：{score} 分，{issues_n} 项问题，结论 {status}").bold = True
        if desc:
            document.add_paragraph(f"审查内容：{desc}")

    document.add_heading("二、技术标五维打分", level=1)
    for dim in dimensions or []:
        document.add_paragraph(
            f"{dim.get('name') or '—'}：{dim.get('score', '—')} 分 / 权重 {dim.get('weight', '—')}%"
        )

    document.add_heading("三、预审问题清单", level=1)
    if not issues:
        document.add_paragraph("本轮无预审问题。")
    for i, issue in enumerate(issues, 1):
        document.add_heading(
            f"{i}. [{issue.get('severity') or '—'}] {issue.get('rule') or '未标注规则'}",
            level=2,
        )
        document.add_paragraph(f"层级：{issue.get('level') or '—'}　　定位：{issue.get('location') or '—'}")
        excerpt = (issue.get("excerpt") or "").strip() or "（未摘录到投标书原文）"
        quote = (issue.get("tenderQuote") or issue.get("tender_quote") or "").strip() or "（未对照到招标条款原文）"
        suggestion = (issue.get("suggestion") or "").strip() or "（无改写建议）"
        document.add_paragraph(f"投标书原文：{excerpt}")
        document.add_paragraph(f"招标要求原文：{quote}")
        document.add_paragraph(f"修改建议：{suggestion}")

    document.add_heading("四、预审结论", level=1)
    document.add_paragraph(
        f"本轮预审综合得分 {overall} 分，风险灯为「{light or '—'}」。"
        f"共发现废标 {waste} 项、扣分 {risk} 项、建议 {suggest} 项。"
        "本报告由系统基于本项目招标文件与投标文件原文自动生成，供投标前自查参考。"
    )

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
