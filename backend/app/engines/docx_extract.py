"""基于 python-docx 的公共文档抽取工具，供五大引擎共用。

本模块是最底层的抽取工具，不反向依赖 tender_form/tender_toc 等上层引擎模块
（它们本身会 import 本模块），避免循环引用。
"""

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _first_run_style(p) -> tuple[str, float, bool]:
    """取段落第一个非空 run 的字体/字号/加粗；run 上没有则回落到段落样式。"""
    style_font = None
    try:
        style_font = p.style.font if p.style is not None else None
    except Exception:
        style_font = None

    def from_run(run) -> tuple[str, float, bool]:
        name = ""
        rPr = run._element.rPr
        if rPr is not None:
            rf = rPr.find(qn("w:rFonts"))
            if rf is not None:
                for key in (qn("w:eastAsia"), qn("w:ascii"), qn("w:hAnsi"), qn("w:cs")):
                    val = rf.get(key)
                    if val:
                        name = val
                        break
        if not name and run.font.name:
            name = run.font.name
        if not name and style_font is not None and style_font.name:
            name = style_font.name
        size = 0.0
        if run.font.size:
            size = round(float(run.font.size.pt), 1)
        elif rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None and sz.get(qn("w:val")):
                try:
                    size = int(sz.get(qn("w:val"))) / 2
                except (TypeError, ValueError):
                    size = 0.0
        if not size and style_font is not None and style_font.size:
            size = round(float(style_font.size.pt), 1)
        bold = bool(run.bold) if run.bold is not None else bool(style_font.bold) if style_font is not None else False
        return name or "宋体", size or 12.0, bold

    for run in p.runs:
        if (run.text or "").strip():
            return from_run(run)
    if style_font is not None:
        name = style_font.name or "宋体"
        size = round(float(style_font.size.pt), 1) if style_font.size else 12.0
        return name, size, bool(style_font.bold)
    return "宋体", 12.0, False


def _paragraph_align(p) -> str:
    """段落对齐方式：""（默认左对齐，不返回）/center/right/justify。"""
    val = ""
    try:
        if p._element.pPr is not None:
            jc = p._element.pPr.find(qn("w:jc"))
            if jc is not None:
                val = (jc.get(qn("w:val")) or "").lower()
    except Exception:
        val = ""
    if p.alignment == WD_ALIGN_PARAGRAPH.CENTER or val in ("center", "middle"):
        return "center"
    if p.alignment == WD_ALIGN_PARAGRAPH.RIGHT or val in ("right", "end"):
        return "right"
    if p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY or val == "both":
        return "justify"
    return ""


def extract_paragraphs(path: str) -> list[dict]:
    """返回非空段落列表：[{index, text, style, outline_level, align}]。"""
    document = docx.Document(path)
    result = []
    for idx, p in enumerate(document.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        style_name = p.style.name if p.style is not None else ""
        try:
            outline_level = p.paragraph_format.outline_level
        except Exception:
            outline_level = None
        font, size_pt, bold = _first_run_style(p)
        result.append(
            {
                "index": idx,
                "text": text,
                "style": style_name,
                "outline_level": outline_level,
                "align": _paragraph_align(p),
                "font": font,
                "fontSizePt": size_pt,
                "bold": bold,
            }
        )
    return result


def extract_full_text(path: str) -> str:
    paragraphs = extract_paragraphs(path)
    return "\n".join(p["text"] for p in paragraphs)


def extract_document_plain_text(path: str) -> str:
    """按正文顺序抽出段落和表格文字，供目录生成阅读整篇招标书。"""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = docx.Document(path)
    chunks: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, document).text.strip()
            if text:
                chunks.append(text)
        elif child.tag == qn("w:tbl"):
            table = Table(child, document)
            for row in table.rows:
                cells: list[str] = []
                for cell in row.cells:
                    cell_text = " ".join(cell.text.split())
                    if cell_text:
                        cells.append(cell_text)
                if cells:
                    chunks.append(" | ".join(cells))
    return "\n".join(chunks)


def has_toc_field(path: str) -> bool:
    document = docx.Document(path)
    xml = document.element.xml
    return "TOC" in xml and ("fldSimple" in xml or "instrText" in xml)


def has_revision_marks(path: str) -> bool:
    document = docx.Document(path)
    xml = document.element.xml
    return "<w:ins " in xml or "<w:del " in xml or "<w:ins>" in xml or "<w:del>" in xml


def get_core_author(path: str) -> str | None:
    document = docx.Document(path)
    props = document.core_properties
    return props.author or props.last_modified_by


def has_comments(path: str) -> bool:
    """检测批注残留（commentRange 或 comments 部件）。"""
    document = docx.Document(path)
    xml = document.element.xml
    if "commentRangeStart" in xml or "commentRangeEnd" in xml or "w:commentReference" in xml:
        return True
    try:
        for rel in document.part.rels.values():
            reltype = (getattr(rel, "reltype", None) or "").lower()
            if "comments" in reltype:
                return True
    except Exception:
        pass
    return False


def has_blank_page_hint(path: str) -> bool:
    """连续空段落达到一页量级时视为空白页线索，不做跨页精确分页。"""
    document = docx.Document(path)
    empty_run = 0
    for p in document.paragraphs:
        if not p.text.strip():
            empty_run += 1
            if empty_run >= 15:
                return True
        else:
            empty_run = 0
    return False
