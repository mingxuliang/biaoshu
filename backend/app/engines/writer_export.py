"""撰写工作台「导出 Word」：把逐章生成/编辑的 markdown-lite 文本还原为真实 .docx。

与前端 ChapterEditor 的行级/行内约定保持一致：
`### ` 三级标题、`## ` 二级标题、`- `/`* ` 项目符号列表（要求标记后有空格，避免与行内 `*斜体*` 冲突）、
`数字)` 或 `（一）` 编号段落，其余为普通段落；行内 `**加粗**` / `*斜体*` / `__下划线__` 会被拆分为独立样式的 run。
`>>> ` 居中、`>> ` 右对齐；连续 `| a | b |` 行还原为表格。
`![说明](/api/writer-images/{id}/file)` 会按本地落盘文件嵌入图片。

不是预审引擎（不产生 Finding），仅供 routers/writer.py 复用。
"""

import io
import os
import re

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Inches, RGBColor

_INLINE_SPLIT_RE = re.compile(r"(\*\*.+?\*\*|__.+?__|\*.+?\*)")
_FONT_WRAP_RE = re.compile(r"\{\{([^|:{}]+)[:|]([^}]+)\}\}(.*?)\{\{/\}\}", re.S)
_CELL_SPAN_RE = re.compile(r"^(?:#c(\d+)#)?(?:#r(\d+)#)?")
_BULLET_RE = re.compile(r"^[-*]\s+")
_NUMBERED_RE = re.compile(r"^\d+[）.)\]].*")
_CN_NUMBERED_RE = re.compile(r"^（[一二三四五六七八九十]+）")
_IMAGE_MD_RE = re.compile(r"^!\[([^\]]*)\]\((/api/writer-images/([^/]+)/file)\)$")
_TABLE_SEP_RE = re.compile(r"^\|?\s*:?-{2,}:?\s*(\|\s*:?-{2,}:?\s*)+\|?\s*$")


def _set_run_font(run, name: str | None, size: str | None) -> None:
    if name:
        run.font.name = name
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn("w:ascii"), name)
        rFonts.set(qn("w:hAnsi"), name)
        rFonts.set(qn("w:eastAsia"), name)
        rFonts.set(qn("w:cs"), name)
    if size:
        m = re.match(r"([0-9.]+)", str(size).strip())
        if m:
            run.font.size = Pt(float(m.group(1)))


def _add_inner_runs(paragraph, text: str, font: str | None = None, size: str | None = None) -> None:
    for part in _INLINE_SPLIT_RE.split(text or ""):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.underline = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        _set_run_font(run, font, size)


def _add_runs_with_inline_styles(paragraph, text: str) -> None:
    raw = text or ""
    pos = 0
    for m in _FONT_WRAP_RE.finditer(raw):
        if m.start() > pos:
            _add_inner_runs(paragraph, raw[pos : m.start()])
        _add_inner_runs(paragraph, m.group(3), m.group(1).strip(), m.group(2).strip())
        pos = m.end()
    if pos == 0:
        _add_inner_runs(paragraph, raw)
    elif pos < len(raw):
        _add_inner_runs(paragraph, raw[pos:])


def _split_align(line: str) -> tuple[object | None, str]:
    raw = line or ""
    stripped = raw.lstrip()
    if stripped.startswith(">>> "):
        return WD_ALIGN_PARAGRAPH.CENTER, stripped[4:]
    if stripped.startswith(">>>"):
        return WD_ALIGN_PARAGRAPH.CENTER, stripped[3:]
    if stripped.startswith(">> "):
        return WD_ALIGN_PARAGRAPH.RIGHT, stripped[3:]
    return None, raw


def _is_table_row(line: str) -> bool:
    t = (line or "").strip()
    return t.startswith("|") and t.count("|") >= 2


def _is_table_sep(line: str) -> bool:
    return bool(_TABLE_SEP_RE.match((line or "").strip()))


def _parse_table_row(line: str) -> list[str]:
    t = (line or "").strip()
    if t.startswith("|"):
        t = t[1:]
    if t.endswith("|") and not t.endswith("\\|"):
        t = t[:-1]
    cells: list[str] = []
    buf: list[str] = []
    brace = 0
    escape = False
    i = 0
    while i < len(t):
        ch = t[i]
        if escape:
            buf.append(ch)
            escape = False
            i += 1
            continue
        if ch == "\\":
            escape = True
            i += 1
            continue
        if ch == "{" and i + 1 < len(t) and t[i + 1] == "{":
            brace += 1
            buf.append("{{")
            i += 2
            continue
        if ch == "}" and i + 1 < len(t) and t[i + 1] == "}":
            brace = max(0, brace - 1)
            buf.append("}}")
            i += 2
            continue
        if ch == "|" and brace == 0:
            cells.append("".join(buf).strip())
            buf = []
            i += 1
            continue
        buf.append(ch)
        i += 1
    cells.append("".join(buf).strip())
    return cells


def _parse_cell(raw: str) -> dict:
    text = raw or ""
    col_span = 1
    row_span = 1
    m = _CELL_SPAN_RE.match(text)
    if m and (m.group(1) or m.group(2)):
        if m.group(1):
            col_span = max(1, int(m.group(1)))
        if m.group(2):
            row_span = max(1, int(m.group(2)))
        text = text[m.end() :]
    return {"text": text, "colspan": col_span, "rowspan": row_span}


def _fill_cell(cell, text: str) -> None:
    cell.text = ""
    parts = re.split(r"<br\s*/?>", text or "", flags=re.I)
    for i, part in enumerate(parts):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        _add_runs_with_inline_styles(p, part)


def _add_markdown_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    parsed = [[_parse_cell(c) for c in row] for row in rows]
    cols = max((sum(c["colspan"] for c in row) for row in parsed), default=1)
    table = document.add_table(rows=len(parsed), cols=max(cols, 1))
    table.style = "Table Grid"
    occupied = [[False] * cols for _ in parsed]
    for i, row in enumerate(parsed):
        c = 0
        for info in row:
            while c < cols and occupied[i][c]:
                c += 1
            if c >= cols:
                break
            cs = max(1, int(info["colspan"]))
            rs = max(1, int(info["rowspan"]))
            cell = table.cell(i, c)
            if cs > 1 or rs > 1:
                end_r = min(i + rs - 1, len(parsed) - 1)
                end_c = min(c + cs - 1, cols - 1)
                if end_r > i or end_c > c:
                    cell.merge(table.cell(end_r, end_c))
            _fill_cell(cell, info["text"])
            for rr in range(i, min(i + rs, len(parsed))):
                for cc in range(c, min(c + cs, cols)):
                    occupied[rr][cc] = True
            c += cs


def render_markdown_block(
    document, line: str, image_paths: dict[str, str] | None = None, layout: dict | None = None
) -> None:
    """把编辑器里的一行 markdown-lite 文本追加为 docx 段落/标题。"""
    align, body = _split_align(line)
    trimmed = body.strip()
    if not trimmed:
        document.add_paragraph("")
        return

    img_match = _IMAGE_MD_RE.match(trimmed)
    if img_match:
        image_id = img_match.group(3)
        path = (image_paths or {}).get(image_id)
        if path and os.path.exists(path):
            try:
                document.add_picture(path, width=Inches(5.4))
            except Exception:  # noqa: BLE001 —— 坏图不阻断整份导出
                document.add_paragraph(f"[图片无法嵌入：{img_match.group(1) or image_id}]")
        else:
            document.add_paragraph(f"[图片缺失：{img_match.group(1) or image_id}]")
        return

    if trimmed.startswith("### "):
        p = document.add_heading(_unwrap_inline(trimmed[4:]), level=3)
        _restyle_heading(p, layout)
        if align is not None:
            p.alignment = align
        return
    if trimmed.startswith("## "):
        p = document.add_heading(_unwrap_inline(trimmed[3:]), level=2)
        _restyle_heading(p, layout)
        if align is not None:
            p.alignment = align
        return

    bullet_match = _BULLET_RE.match(trimmed)
    if bullet_match:
        p = document.add_paragraph(style="List Bullet")
        _add_runs_with_inline_styles(p, trimmed[bullet_match.end() :].strip())
        if align is not None:
            p.alignment = align
        return

    if _NUMBERED_RE.match(trimmed) or _CN_NUMBERED_RE.match(trimmed):
        p = document.add_paragraph()
        _add_runs_with_inline_styles(p, trimmed)
        if align is not None:
            p.alignment = align
        return

    p = document.add_paragraph()
    _add_runs_with_inline_styles(p, trimmed)
    _apply_body_indent(p, trimmed, layout)
    if align is not None:
        p.alignment = align


def _node_level(outline: list[dict], node_id: str, _depth: int = 0) -> int:
    node = next((n for n in outline if n.get("id") == node_id), None)
    if not node or not node.get("parentId"):
        return _depth
    return _node_level(outline, node["parentId"], _depth + 1)


_FONT_PT = {"小五": 9, "五号": 10.5, "小四": 12, "四号": 14, "小三": 15, "三号": 16, "小二": 18}
_LINE_SPACING = {
    "1.5倍行距": 1.5,
    "2倍行距": 2.0,
    "固定值28磅": 28,
    "固定值30磅": 30,
}


def _unwrap_inline(text: str) -> str:
    t = _FONT_WRAP_RE.sub(lambda m: m.group(3), text or "")
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", t)
    t = re.sub(r"__(.+?)__", r"\1", t)
    t = re.sub(r"(?<!\*)\*(.+?)\*(?!\*)", r"\1", t)
    return t.strip()


def _set_style_font(style, name: str, size_pt: float) -> None:
    font = style.font
    font.name = name
    font.size = Pt(size_pt)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:cs"), name)


def _layout_body_pt(layout: dict | None) -> float:
    if not layout:
        return 12.0
    try:
        if layout.get("bodySizePt") is not None:
            return float(layout["bodySizePt"])
    except (TypeError, ValueError):
        pass
    return float(_FONT_PT.get(str(layout.get("fontSize") or ""), 12))


def _restyle_heading(paragraph, layout: dict | None) -> None:
    if not layout:
        return
    hfont = str(layout.get("headingFont") or layout.get("bodyFont") or "宋体")
    try:
        hpt = float(layout.get("headingSizePt") or _layout_body_pt(layout))
    except (TypeError, ValueError):
        hpt = _layout_body_pt(layout)
    hbold = bool(layout.get("headingBold", True))
    paragraph.paragraph_format.first_line_indent = Pt(0)
    for run in paragraph.runs:
        _set_run_font(run, hfont, f"{hpt:g}pt")
        run.bold = hbold
        run.font.color.rgb = RGBColor(0, 0, 0)


def _apply_body_indent(paragraph, text: str, layout: dict | None) -> None:
    if not layout:
        return
    try:
        indent_pt = float(layout.get("indentPt") or 0)
    except (TypeError, ValueError):
        return
    if indent_pt < 8:
        return
    stripped = (text or "").lstrip()
    if stripped.startswith("　") or stripped.startswith(" "):
        return
    paragraph.paragraph_format.first_line_indent = Pt(indent_pt)


def _apply_layout(document, layout: dict | None) -> None:
    if not layout:
        return
    section = document.sections[0]
    margins = layout.get("margins") or {}
    if isinstance(margins, dict):
        if margins.get("top") is not None:
            section.top_margin = Cm(float(margins["top"]))
        if margins.get("bottom") is not None:
            section.bottom_margin = Cm(float(margins["bottom"]))
        if margins.get("left") is not None:
            section.left_margin = Cm(float(margins["left"]))
        if margins.get("right") is not None:
            section.right_margin = Cm(float(margins["right"]))

    body_font = str(layout.get("bodyFont") or "宋体")
    body_pt = _layout_body_pt(layout)
    _set_style_font(document.styles["Normal"], body_font, body_pt)

    heading_font = str(layout.get("headingFont") or body_font)
    try:
        heading_pt = float(layout.get("headingSizePt") or body_pt)
    except (TypeError, ValueError):
        heading_pt = body_pt
    heading_bold = bool(layout.get("headingBold", True))
    for i in range(1, 4):
        try:
            hstyle = document.styles[f"Heading {i}"]
        except KeyError:
            continue
        _set_style_font(hstyle, heading_font, heading_pt)
        hstyle.font.bold = heading_bold
        hstyle.font.color.rgb = RGBColor(0, 0, 0)
        hstyle.paragraph_format.first_line_indent = Pt(0)

    pf = document.styles["Normal"].paragraph_format
    spacing = layout.get("lineSpacing")
    try:
        mul = float(layout["lineSpacingMul"]) if layout.get("lineSpacingMul") is not None else None
    except (TypeError, ValueError):
        mul = None
    multiple = mul if mul else _LINE_SPACING.get(str(spacing or ""), 1.5)
    if spacing in ("固定值28磅", "固定值30磅") and mul is None:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(multiple)
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = multiple


def chapters_to_docx(
    project_name: str,
    outline: list[dict],
    chapter_contents: dict[str, str],
    image_paths: dict[str, str] | None = None,
    layout: dict | None = None,
) -> bytes:
    """按目录顺序把全部章节内容拼接为一份完整 .docx。"""
    from .e_writer import chapter_kind

    document = docx.Document()
    _apply_layout(document, layout)
    if project_name:
        document.add_heading(project_name, level=0)

    for node in outline:
        level = min(_node_level(outline, node["id"]) + 1, 3)
        heading_text = f"{node.get('num', '')} {node.get('title', '')}".strip()
        document.add_heading(heading_text, level=level)

        kind = chapter_kind(
            str(node.get("title") or ""),
            node.get("part"),
            str(node.get("idea") or ""),
            str(node.get("requirement") or ""),
        )
        body_layout = layout if kind == "tech" else None

        content = chapter_contents.get(node["id"], "")
        if not content.strip():
            continue
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            if _is_table_row(lines[i]) or _is_table_sep(lines[i]):
                rows: list[list[str]] = []
                while i < len(lines) and (_is_table_row(lines[i]) or _is_table_sep(lines[i])):
                    if not _is_table_sep(lines[i]):
                        rows.append(_parse_table_row(lines[i]))
                    i += 1
                _add_markdown_table(document, rows)
                continue
            render_markdown_block(document, lines[i], image_paths, body_layout)
            i += 1

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
