"""修改闭环（Review）用的辅助构建函数：把 docx 扁平段落分组为章节树，
把 ReviewFinding 锚定回具体段落，并把编辑器序列化出的 blocks 重新写回 docx。

不是预审引擎（不产生 Finding），仅供 routers/revision.py 复用。
"""

import difflib
import io
import re

import docx

_PARA_INDEX_RE = re.compile(r"段落\s*(\d+)")
_LCS_RATIO_THRESHOLD = 0.5
_WS_RE = re.compile(r"\s+")
_CHAPTER = re.compile(r"^第[0-9一二三四五六七八九十百零]+[章节篇]")
_CN_DOT = re.compile(r"^([一二三四五六七八九十]+)、")
_CN_PAREN = re.compile(r"^[（(]([一二三四五六七八九十]+)[）)]")
_DOTTED = re.compile(r"^(\d+\.\d+(?:\.\d+)*)")
_SEQ = re.compile(r"^(\d{1,2})[.．、](?!\d)")
_ATTACH = re.compile(r"^附件[0-9一二三四五六七八九十]")


def _normalize(s: str) -> str:
    return _WS_RE.sub("", s or "")


def _heading_level(p: dict, body_size_pt: float = 12.0) -> int | None:
    """判断该段落是否为标题，返回 1/2/3；非标题返回 None。

    投标书大量标题并不使用 Word「标题 1/2」样式，只是「第一章 / 一、 / 1.1」编号
    或加粗放大字号。只认样式名会把一级二级标题全部当成正文。
    """
    style_name = p.get("style") or ""
    m = re.search(r"(?:Heading|标题)\s*(\d+)", style_name, re.IGNORECASE)
    if m:
        return min(max(int(m.group(1)), 1), 3)
    outline_level = p.get("outline_level")
    if outline_level is not None and outline_level <= 2:
        return outline_level + 1

    text = (p.get("text") or "").strip()
    if not text:
        return None
    if _CHAPTER.match(text) or _ATTACH.match(text):
        return 1
    if _CN_DOT.match(text) and len(text) <= 48 and "。" not in text:
        return 1
    if _CN_PAREN.match(text) and len(text) <= 48 and "。" not in text:
        return 2
    dotted = _DOTTED.match(text)
    if dotted and len(text) <= 60 and "。" not in text:
        return min(1 + dotted.group(1).count("."), 3)

    try:
        size_pt = float(p.get("fontSizePt") or 0)
    except (TypeError, ValueError):
        size_pt = 0.0
    bold = bool(p.get("bold"))
    looks_title = bold or size_pt >= body_size_pt + 1.5
    # 「1. 」编号正文很常见，不能单凭序号升成一级标题；必须加粗或明显大于正文字号。
    seq = _SEQ.match(text)
    if seq and looks_title and len(text) <= 40 and "。" not in text:
        try:
            if int(seq.group(1)) <= 40:
                return 1
        except ValueError:
            pass

    if len(text) <= 32 and "。" not in text and looks_title:
        if size_pt >= body_size_pt + 4:
            return 1
        return 2
    return None


def build_sections(paragraphs: list[dict]) -> list[dict]:
    """把 extract_paragraphs() 的扁平段落列表分组为 BidSection[]。

    非标题段落挂到最近一个标题所在的 section；文档开头若无标题，
    生成一个默认的「文档开头」占位 section 承接。
    """
    body_sizes = [
        float(p.get("fontSizePt") or 12)
        for p in paragraphs
        if not p.get("bold") and len(p.get("text") or "") > 20
    ]
    body_size = sorted(body_sizes)[len(body_sizes) // 2] if body_sizes else 12.0

    sections: list[dict] = []
    current: dict | None = None
    sec_seq = 0
    para_seq = 0

    for p in paragraphs:
        level = _heading_level(p, body_size)
        if level is not None:
            sec_seq += 1
            current = {
                "id": f"sec-{sec_seq}",
                "heading": p["text"],
                "level": level,
                "paragraphs": [],
                "align": p.get("align") or "",
                "font": p.get("font") or "",
                "fontSizePt": p.get("fontSizePt"),
                "bold": bool(p.get("bold")),
            }
            sections.append(current)
            continue

        if current is None:
            sec_seq += 1
            current = {
                "id": f"sec-{sec_seq}",
                "heading": "文档开头",
                "level": 1,
                "paragraphs": [],
                "align": "",
            }
            sections.append(current)

        para_seq += 1
        current["paragraphs"].append(
            {
                "id": f"p-{para_seq}",
                "text": p["text"],
                "index": p["index"],
                "align": p.get("align") or "",
                "font": p.get("font") or "",
                "fontSizePt": p.get("fontSizePt"),
                "bold": bool(p.get("bold")),
            }
        )

    return sections


def _longest_common_span(excerpt: str, text: str) -> tuple[int, int, float]:
    """返回 (text 里的起点, 长度, 相对 excerpt 长度的命中比例)。"""
    if not excerpt or not text:
        return 0, 0, 0.0
    matcher = difflib.SequenceMatcher(None, excerpt, text, autojunk=False)
    match = matcher.find_longest_match(0, len(excerpt), 0, len(text))
    ratio = match.size / max(len(excerpt), 1)
    return match.b, match.size, ratio


def _find_target(excerpt: str, all_paragraphs: list[dict], used_para_ids: set[str]) -> tuple[dict | None, str]:
    """依次尝试「原文完全包含 excerpt」「归一化包含」「最长公共子串」三档匹配，
    返回 (命中的段落, 用于前端高亮的 highlight 文本——必须是该段落 text 的真子串)。
    """
    if not excerpt:
        return None, ""

    for para in all_paragraphs:
        if para["id"] in used_para_ids:
            continue
        if excerpt in para["text"]:
            return para, excerpt

    norm_excerpt = _normalize(excerpt)
    if norm_excerpt:
        candidates = []
        for para in all_paragraphs:
            if para["id"] in used_para_ids:
                continue
            norm_text = _normalize(para["text"])
            if not norm_text:
                continue
            if norm_excerpt in norm_text or (len(norm_excerpt) > 8 and norm_text in norm_excerpt):
                candidates.append(para)
        if candidates:
            best = min(candidates, key=lambda p: abs(len(p["text"]) - len(excerpt)))
            start, size, _ = _longest_common_span(excerpt, best["text"])
            highlight = best["text"][start : start + size] if size >= 4 else best["text"]
            return best, highlight

    best_para = None
    best_ratio = 0.0
    best_span = (0, 0)
    for para in all_paragraphs:
        if para["id"] in used_para_ids:
            continue
        start, size, ratio = _longest_common_span(excerpt, para["text"])
        if size < 6:
            continue
        if ratio > best_ratio:
            best_ratio = ratio
            best_para = para
            best_span = (start, size)
    if best_para is not None and best_ratio >= _LCS_RATIO_THRESHOLD:
        start, size = best_span
        highlight = best_para["text"][start : start + size]
        return best_para, highlight

    return None, ""


def anchor_findings(sections: list[dict], issues: list[dict]) -> list[dict]:
    """把 issues（对齐 PreReviewIssueOut 字段）尽量锚定到 sections 里的具体段落上，
    写入 paragraph["problem"] = {issueId, highlight}。每个段落最多锚定一个问题
    （与前端 BidParagraph.problem 的单问题结构保持一致，先匹配先占用）。

    location 里带「段落 N」的优先按段号精确定位；否则用 excerpt 做「原文包含 /
    归一化包含 / 最长公共子串」三档兜底匹配——用最长公共子串而不是整串相似度，
    是因为 excerpt 通常只是段落里的一小段引用，整串 ratio 会被段落长度稀释到
    很低，导致大量真实可定位的问题被判定为「无法定位」。全篇级问题（既没有
    段落号，excerpt 也找不到匹配）保持不锚定，交给前端提示用户无法跳转。
    """
    all_paragraphs = [para for sec in sections for para in sec["paragraphs"]]
    used_para_ids: set[str] = set()

    for issue in issues:
        location = issue.get("location", "")
        excerpt = (issue.get("excerpt") or "").strip()
        target = None
        highlight = ""

        m = _PARA_INDEX_RE.search(location)
        if m:
            idx = int(m.group(1))
            for para in all_paragraphs:
                if para["index"] == idx and para["id"] not in used_para_ids:
                    target = para
                    highlight = excerpt if excerpt and excerpt in para["text"] else para["text"]
                    break

        if target is None:
            target, highlight = _find_target(excerpt, all_paragraphs, used_para_ids)

        if target is None:
            continue

        target["problem"] = {"issueId": issue["id"], "highlight": highlight or target["text"]}
        used_para_ids.add(target["id"])

    for sec in sections:
        for para in sec["paragraphs"]:
            para.pop("index", None)

    return sections


def blocks_to_docx(blocks: list[dict]) -> bytes:
    """把编辑器序列化出的 {type: heading|paragraph, level?, text} 顺序块重新写为 .docx。

    仅还原标题层级与纯文本段落，不还原粗体/表格/链接等 Lexical 高级样式。
    优先使用 writeback_docx() 在原稿上改字，以保留字体、字号与段落样式。
    """
    document = docx.Document()
    for block in blocks:
        text = block.get("text", "")
        if block.get("type") == "heading":
            level = min(max(int(block.get("level") or 1), 1), 3)
            document.add_heading(text, level=level)
        else:
            document.add_paragraph(text)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _replace_para_text(para, text: str) -> None:
    """只改文字，保留首个 run 的字体/加粗/字号；其余 run 清空以免重复拼接。"""
    if (para.text or "").strip() == text:
        return
    runs = para.runs
    if not runs:
        para.add_run(text)
        return
    runs[0].text = text
    for run in runs[1:]:
        run.text = ""


def writeback_docx(original_path: str, blocks: list[dict]) -> bytes:
    """在原始投标书 docx 上按段落对齐回写编辑器文本，尽量保留原稿样式。

    用 SequenceMatcher 把非空原稿段落与编辑器非空块对齐；只处理 equal/replace
    中一一对应的段落（改字不改结构）。对齐质量过低时退回 blocks_to_docx，
    避免把完全对不上的稿硬塞进原稿。
    """
    from pathlib import Path

    if not blocks:
        return Path(original_path).read_bytes()

    try:
        document = docx.Document(original_path)
    except Exception:
        return blocks_to_docx(blocks)

    orig_paras = [p for p in document.paragraphs if (p.text or "").strip()]
    new_texts = [(b.get("text") or "").strip() for b in blocks if (b.get("text") or "").strip()]
    if not orig_paras or not new_texts:
        return Path(original_path).read_bytes()

    orig_texts = [p.text.strip() for p in orig_paras]
    joined_ratio = difflib.SequenceMatcher(
        None, "\n".join(orig_texts), "\n".join(new_texts), autojunk=False
    ).ratio()
    if joined_ratio < 0.25:
        return blocks_to_docx(blocks)

    matcher = difflib.SequenceMatcher(a=orig_texts, b=new_texts, autojunk=False)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag != "replace":
            continue
        n = min(i2 - i1, j2 - j1)
        for k in range(n):
            _replace_para_text(orig_paras[i1 + k], new_texts[j1 + k])

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
