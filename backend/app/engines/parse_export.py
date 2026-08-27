"""把评标尺子导出为可读的解析报告 Word。按固定一级/二级指标输出，空字段也保留标题。"""

import io

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .parse_schema import merge_tree


def checklist_to_docx(
    *,
    project_name: str,
    project_code: str,
    version: int,
    locked: bool,
    data: dict,
) -> bytes:
    document = docx.Document()
    title = document.add_heading("招标文件解析报告", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = document.add_paragraph()
    meta.add_run(f"项目：{project_name or '（未命名）'}\n")
    meta.add_run(f"招标编号：{project_code or '—'}\n")
    meta.add_run(f"评标尺子版本：v{version}（{'已锁定' if locked else '草稿'}）")

    for dim in merge_tree(data.get("dimensions")):
        document.add_heading(dim["label"], level=1)
        for item in dim["items"]:
            document.add_heading(item["label"], level=2)
            for sec in item["sections"]:
                document.add_heading(sec["title"], level=3)
                for row in sec["rows"]:
                    p = document.add_paragraph()
                    p.add_run(f"{row['label']}：").bold = True
                    content = (row.get("content") or "").strip()
                    p.add_run(content if content else "（未从招标文件中抽取到该项内容）")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
