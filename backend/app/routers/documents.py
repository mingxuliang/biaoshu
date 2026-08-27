import io
import os

import docx
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import Response
from sqlalchemy.orm import Session

from ..auth import get_current_user
from ..config import get_settings
from ..db import get_db
from ..models import BidDocument, User
from ..permissions import PERM_REVIEW, PERM_WRITER, require_any_perm, require_project
from ..schemas import BidDocumentSummaryOut, UploadDocOut
from .. import storage

router = APIRouter(prefix="/api", tags=["documents"])


def _bid_doc_to_summary(d: BidDocument) -> BidDocumentSummaryOut:
    return BidDocumentSummaryOut(
        id=d.id,
        filename=d.filename,
        source=d.source,
        sizeBytes=d.size_bytes,
        uploadedAt=d.uploaded_at.isoformat(),
    )


@router.post("/bid-documents", response_model=UploadDocOut)
async def upload_bid_document(
    project_id: str = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> UploadDocOut:
    require_project(db, current_user, project_id)
    require_any_perm(current_user, PERM_WRITER, PERM_REVIEW)
    filename = file.filename or ""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".doc":
        raise HTTPException(400, "暂不支持旧版 .doc 格式，请在 Word 中另存为 .docx 后重新上传")
    if ext != ".docx":
        raise HTTPException(400, "仅支持 .docx 格式的 Word 文档")

    content = await file.read()
    try:
        docx.Document(io.BytesIO(content))
    except Exception as exc:
        raise HTTPException(400, "文档已损坏或无法解析，请重新上传") from exc

    key = storage.put_bytes(f"bid-documents/{project_id}", content, ext)
    doc = BidDocument(
        project_id=project_id,
        filename=filename,
        storage_path=key,
        size_bytes=len(content),
        source="upload",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return UploadDocOut(id=doc.id, filename=doc.filename, size_bytes=doc.size_bytes, source=doc.source)


@router.get("/projects/{project_id}/bid-documents", response_model=list[BidDocumentSummaryOut])
def list_project_bid_documents(
    project_id: str, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
) -> list[BidDocumentSummaryOut]:
    require_project(db, current_user, project_id)
    docs = (
        db.query(BidDocument)
        .filter(BidDocument.project_id == project_id)
        .order_by(BidDocument.uploaded_at.desc())
        .all()
    )
    return [_bid_doc_to_summary(d) for d in docs]


@router.get("/bid-documents/{doc_id}/download")
def download_bid_document(
    doc_id: str, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
) -> Response:
    doc = db.get(BidDocument, doc_id)
    if not doc or not storage.exists(doc.storage_path):
        raise HTTPException(404, "文件不存在")
    require_project(db, current_user, doc.project_id)
    try:
        return storage.http_response(doc.storage_path, filename=doc.filename)
    except FileNotFoundError:
        raise HTTPException(404, "文件不存在")


@router.post("/bid-documents/from-sample", response_model=UploadDocOut)
def use_sample_document(
    project_id: str = Form(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
) -> UploadDocOut:
    """基于内置示例投标书快速创建一份可供预审的文档（供测试用）。"""
    require_project(db, current_user, project_id, PERM_WRITER)
    settings = get_settings()
    sample_path = os.path.join(settings.sample_dir, "demo_bid.docx")
    if not os.path.exists(sample_path):
        raise HTTPException(500, "示例文档缺失，请检查 backend/sample_data/demo_bid.docx 是否存在")

    with open(sample_path, "rb") as fh:
        content = fh.read()
    key = storage.put_bytes(f"bid-documents/{project_id}", content, ".docx")

    doc = BidDocument(
        project_id=project_id,
        filename="demo_bid.docx（工作台示例文档）",
        storage_path=key,
        size_bytes=len(content),
        source="workbench",
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)

    return UploadDocOut(id=doc.id, filename=doc.filename, size_bytes=doc.size_bytes, source=doc.source)
