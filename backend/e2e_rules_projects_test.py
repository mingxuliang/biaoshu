"""端到端验证：预审规则真实后端 + 项目中心团队/文件/时间线/招标文件。

1. 登录
2. 规则种子存在；关闭一条虚词；新增并激活权重模板；修改阈值
3. 列出真实用户；新建项目并分配成员；时间线与文件归档按项目隔离
"""

import io
import json
import sys
import urllib.error
import urllib.parse
import urllib.request

import docx

sys.stdout.reconfigure(encoding="utf-8")

BASE = "http://localhost:8000"


def call(method: str, path: str, token: str | None = None, json_body: dict | None = None, data: bytes | None = None, content_type: str | None = None):
    url = f"{BASE}{path}"
    body = data
    if json_body is not None:
        body = json.dumps(json_body).encode("utf-8")
        content_type = "application/json"
    req = urllib.request.Request(url, data=body, method=method)
    if content_type:
        req.add_header("Content-Type", content_type)
    if token:
        req.add_header("Authorization", f"Bearer {token}")
    try:
        with urllib.request.urlopen(req) as resp:
            return resp.status, resp.read()
    except urllib.error.HTTPError as e:
        return e.code, e.read()


def main() -> None:
    status, body = call("POST", "/api/auth/login", json_body={"email": "chen@zhibiaoyun.com", "password": "123456"})
    assert status == 200, f"登录失败: {status} {body}"
    token = json.loads(body)["token"]
    print("login ok")

    status, body = call("GET", "/api/rules/weight-templates", token=token)
    assert status == 200, body
    weights = json.loads(body)
    assert any(t["active"] for t in weights), "应有一条启用中的权重模板"
    print("weight templates:", len(weights), "active:", [t["name"] for t in weights if t["active"]])

    status, body = call("GET", "/api/rules/word-rules", token=token)
    assert status == 200, body
    words = json.loads(body)
    assert len(words) > 10, "虚词种子数量不足"
    target = next((w for w in words if w["word"] == "确保"), words[0])
    status, body = call("PATCH", f"/api/rules/word-rules/{target['id']}", token=token, json_body={"enabled": False, "rewrite": "改为可量化指标"})
    assert status == 200, body
    updated_word = json.loads(body)
    assert updated_word["enabled"] is False
    print("disabled filler word:", updated_word["word"])

    status, body = call(
        "POST",
        "/api/rules/weight-templates",
        token=token,
        json_body={
            "name": "验证用模板",
            "completeness": 40,
            "relevance": 20,
            "compliance": 20,
            "feasibility": 10,
            "standardization": 10,
            "scope": "测试",
        },
    )
    assert status == 200, body
    created_tpl = json.loads(body)
    status, body = call("POST", f"/api/rules/weight-templates/{created_tpl['id']}/activate", token=token)
    assert status == 200, body
    assert json.loads(body)["active"] is True
    print("activated template:", created_tpl["name"])

    status, body = call("GET", "/api/rules/thresholds", token=token)
    assert status == 200, body
    thresholds = json.loads(body)
    assert {t["key"] for t in thresholds} >= {"filler_density_safe", "asset_liability_ratio_max"}
    ratio = next(t for t in thresholds if t["key"] == "asset_liability_ratio_max")
    status, body = call("PATCH", f"/api/rules/thresholds/{ratio['id']}", token=token, json_body={"value": 70})
    assert status == 200, body
    assert json.loads(body)["value"] == 70
    print("updated threshold asset_liability_ratio_max -> 70")

    status, body = call(
        "POST",
        "/api/rules/packages",
        token=token,
        json_body={"name": "测试细则包", "region": "测试区", "items": ["条目A", "条目B"]},
    )
    assert status == 200, body
    print("created rule package")

    status, body = call("GET", "/api/users", token=token)
    assert status == 200, body
    users = json.loads(body)
    if len(users) < 2:
        status, body = call(
            "POST",
            "/api/users",
            token=token,
            json_body={"name": "规则验证成员", "email": "e2e-rules-member@zby.ai", "role": "撰写专家"},
        )
        assert status == 200, body
        status, body = call("GET", "/api/users", token=token)
        assert status == 200, body
        users = json.loads(body)
    assert len(users) >= 2, "应至少有管理员与一名可分配成员"
    print("users:", len(users))

    status, body = call(
        "POST",
        "/api/projects",
        token=token,
        json_body={"name": "规则中心隔离对照项目", "code": "E2E-RULES-OLD", "type": "工程"},
    )
    assert status == 200, body
    old_project = json.loads(body)
    old_id = old_project["id"]

    status, body = call("GET", f"/api/projects/{old_id}/timeline", token=token)
    assert status == 200, body
    timeline_old = json.loads(body)
    assert timeline_old[0]["status"] == "已完成"
    print("old project timeline:", [(s["label"], s["status"]) for s in timeline_old])

    status, body = call(
        "POST",
        "/api/projects",
        token=token,
        json_body={"name": "规则与项目中心验证项目", "code": "E2E-RULES-001", "type": "工程", "owner": "陈立群"},
    )
    assert status == 200, body
    new_project = json.loads(body)
    project_id = new_project["id"]

    member_ids = [users[0]["id"], users[1]["id"]]
    status, body = call("PUT", f"/api/projects/{project_id}/members", token=token, json_body={"user_ids": member_ids})
    assert status == 200, body
    team = json.loads(body)
    assert len(team) == 2
    print("new project team:", [m["name"] for m in team])

    document = docx.Document()
    document.add_heading("验证用招标文件", level=1)
    document.add_paragraph("评标办法：技术标 60 分，商务标 40 分。")
    buf = io.BytesIO()
    document.save(buf)
    docx_bytes = buf.getvalue()

    boundary = "----e2eBoundary"
    filename = "e2e-tender.docx"
    parts = []
    parts.append(f"--{boundary}\r\nContent-Disposition: form-data; name=\"project_id\"\r\n\r\n{project_id}\r\n".encode())
    parts.append(
        (
            f"--{boundary}\r\nContent-Disposition: form-data; name=\"file\"; filename=\"{filename}\"\r\n"
            "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
        ).encode()
        + docx_bytes
        + b"\r\n"
    )
    parts.append(f"--{boundary}--\r\n".encode())
    status, body = call(
        "POST",
        "/api/tender-documents",
        token=token,
        data=b"".join(parts),
        content_type=f"multipart/form-data; boundary={boundary}",
    )
    assert status == 200, f"上传招标文件失败: {status} {body}"
    tender = json.loads(body)
    print("uploaded tender:", tender["id"])

    status, body = call("GET", f"/api/projects/{project_id}/tender-documents", token=token)
    assert status == 200, body
    tenders = json.loads(body)
    assert any(d["id"] == tender["id"] for d in tenders)

    status, body = call("GET", f"/api/projects/{project_id}", token=token)
    assert status == 200, body
    refreshed = json.loads(body)
    assert refreshed.get("tenderDoc"), "项目卡片应回填真实招标文件徽标"
    print("project tenderDoc:", refreshed["tenderDoc"])

    status, body = call("GET", f"/api/projects/{project_id}/documents", token=token)
    assert status == 200, body
    docs = json.loads(body)
    assert len(docs["tenderDocuments"]) >= 1
    status, body = call("GET", f"/api/projects/{old_id}/documents", token=token)
    assert status == 200, body
    docs_other = json.loads(body)
    assert docs["tenderDocuments"][0]["id"] not in [d["id"] for d in docs_other["tenderDocuments"]]
    print("documents isolated across projects")

    status, body = call("GET", f"/api/tender-documents/{tender['id']}/download", token=token)
    assert status == 200, body
    assert len(body) > 100
    print("tender download size:", len(body))

    status, body = call("GET", f"/api/projects/{project_id}/timeline", token=token)
    assert status == 200, body
    timeline_new = json.loads(body)
    parse_stage = next(s for s in timeline_new if s["id"] == "tender_parse")
    assert parse_stage["status"] in ("进行中", "已完成")
    print("new project timeline:", [(s["label"], s["status"]) for s in timeline_new])
    assert timeline_new != timeline_old

    print("\nALL CHECKS PASSED")


if __name__ == "__main__":
    main()
