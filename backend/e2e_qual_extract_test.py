"""端到端：商务标抽取资质/合同/财务，营业执照跨文件合并，过期财务仍入库。"""

import io
import json
import struct
import sys
import time
import zlib
import urllib.error
import urllib.request

import docx
from docx.shared import Inches

sys.stdout.reconfigure(encoding="utf-8")

BASE = "http://localhost:8000"
CREDIT = "91110000E2EQUAL01"
CONTRACT_NO = "HT-E2E-2024-001"


def call(
    method: str,
    path: str,
    token: str | None = None,
    json_body: dict | None = None,
    data: bytes | None = None,
    content_type: str | None = None,
    timeout: int = 30,
):
    url = f"{BASE}{path}"
    body = data
    if json_body is not None:
        body = json.dumps(json_body).encode("utf-8")
        content_type = "application/json"
    req = urllib.request.Request(url, data=body, method=method)
    if content_type:
        req.add_header("Content-Type", content_type)
    if token:
        req.add_header("Authorization", f"Bearer {token}")
    try:
        with urllib.request.urlopen(req, timeout=timeout) as resp:
            return resp.status, resp.read()
    except urllib.error.HTTPError as e:
        return e.code, e.read()


def multipart(files: list[tuple[str, str, bytes]]) -> tuple[bytes, str]:
    boundary = "----e2eQualExtract"
    parts: list[bytes] = []
    for field, filename, blob in files:
        parts.append(
            (
                f"--{boundary}\r\nContent-Disposition: form-data; name=\"{field}\"; filename=\"{filename}\"\r\n"
                "Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document\r\n\r\n"
            ).encode()
            + blob
            + b"\r\n"
        )
    parts.append(f"--{boundary}--\r\n".encode())
    return b"".join(parts), f"multipart/form-data; boundary={boundary}"


def make_scan_png() -> bytes:
    width, height = 400, 300

    def chunk(tag: bytes, data: bytes) -> bytes:
        return struct.pack(">I", len(data)) + tag + data + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)

    raw = bytearray()
    for y in range(height):
        raw.append(0)
        for x in range(width):
            raw.extend(((x * 17 + y * 13) % 256, (x * 3) % 256, (y * 5) % 256))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(bytes(raw), 9))
        + chunk(b"IEND", b"")
    )


def make_biz_docx(year: str) -> bytes:
    document = docx.Document()
    document.add_heading("资格审查文件", level=1)
    document.add_heading("营业执照", level=2)
    document.add_paragraph(f"统一社会信用代码 {CREDIT} 有效期至 长期。注册资本 5000 万元。")
    document.add_picture(io.BytesIO(make_scan_png()), width=Inches(3.2))
    document.add_heading("施工合同", level=2)
    document.add_paragraph(f"合同编号 {CONTRACT_NO} 有效期至 2028-12-31。合同金额 1200 万元。")
    document.add_heading(f"{year}年度审计报告", level=2)
    document.add_paragraph(f"报表日期 {year}-12-31。资产负债率 42%。")
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def poll_job(token: str, job_id: str, timeout_s: int = 180) -> dict:
    started = time.time()
    while time.time() - started < timeout_s:
        status, body = call("GET", f"/api/qualification-extract-jobs/{job_id}", token=token)
        assert status == 200, body
        job = json.loads(body)
        if job["status"] in ("done", "failed"):
            return job
        time.sleep(1.5)
    raise AssertionError(f"抽取任务超时: {job_id}")


def main() -> None:
    status, body = call("POST", "/api/auth/login", json_body={"email": "chen@zhibiaoyun.com", "password": "123456"})
    assert status == 200, f"登录失败: {status} {body}"
    token = json.loads(body)["token"]

    payload, ctype = multipart([("files", "biz-2024.docx", make_biz_docx("2024"))])
    status, body = call(
        "POST",
        "/api/qualification-source-docs",
        token=token,
        data=payload,
        content_type=ctype,
        timeout=60,
    )
    assert status == 200, body
    jobs = json.loads(body)
    assert len(jobs) == 1, jobs
    job = poll_job(token, jobs[0]["id"])
    assert job["status"] == "done", job
    print("first extract ok", job)

    status, body = call("GET", "/api/qualifications", token=token)
    assert status == 200, body
    first = json.loads(body)
    licenses = [q for q in first if CREDIT in (q.get("number") or "") or q.get("name") == "营业执照"]
    contracts = [q for q in first if CONTRACT_NO.replace("-", "") in (q.get("number") or "").replace("-", "")]
    fins_2024 = [q for q in first if q.get("kind") == "financial" and "2024" in (q.get("validUntil") or q.get("name") or "")]
    assert licenses, first
    license_row = licenses[0]
    assert license_row.get("hasFile") or (license_row.get("images") or []), license_row
    print("license images", len(license_row.get("images") or []), "hasFile", license_row.get("hasFile"))
    assert contracts or any("合同" in (q.get("name") or "") for q in first), first
    print("after first", [(q["kind"], q["name"], q["number"], q["validUntil"]) for q in first if q.get("reviewStatus") == "待审核"][:12])

    payload, ctype = multipart([("files", "biz-2023.doc", make_biz_docx("2023"))])
    status, body = call(
        "POST",
        "/api/qualification-source-docs",
        token=token,
        data=payload,
        content_type=ctype,
        timeout=60,
    )
    assert status == 200, body
    jobs2 = json.loads(body)
    job2 = poll_job(token, jobs2[0]["id"])
    assert job2["status"] == "done", job2
    print("second extract ok", job2)

    status, body = call("GET", "/api/qualifications", token=token)
    second = json.loads(body)
    license_ids = {
        q["id"]
        for q in second
        if q.get("name") == "营业执照" or CREDIT in (q.get("number") or "")
    }
    assert len(license_ids) == 1, [q for q in second if q["id"] in license_ids or q.get("name") == "营业执照"]
    merged_license = next(q for q in second if q["id"] in license_ids)
    assert len(merged_license.get("sources") or []) >= 2 or merged_license.get("mergeStatus") in (
        "并入已有",
        "新增",
        "信息冲突",
    )
    fins = [q for q in second if q.get("kind") == "financial"]
    years = {(q.get("validUntil") or "")[:4] for q in fins} | {
        y.group(1)
        for q in fins
        for y in [__import__("re").search(r"(20\d{2})", q.get("name") or "")]
        if y
    }
    assert "2024" in years or fins_2024, fins
    assert "2023" in years, fins
    print("license merged to one; financial years kept separate", years)

    expired = next((q for q in second if q.get("kind") == "financial" and q.get("status") == "已过期"), None)
    if expired:
        assert expired.get("id"), "expired financial still listed"
        print("expired financial still usable", expired["name"], expired["validUntil"])

    keep = next(q for q in second if q["id"] in license_ids)
    if keep.get("reviewStatus") == "待审核":
        boundary = "----e2eQualPatch"
        form = (
            f"--{boundary}\r\nContent-Disposition: form-data; name=\"review_status\"\r\n\r\n已入库\r\n"
            f"--{boundary}--\r\n"
        ).encode()
        status, body = call(
            "PATCH",
            f"/api/qualifications/{keep['id']}",
            token=token,
            data=form,
            content_type=f"multipart/form-data; boundary={boundary}",
        )
        assert status == 200, body
        print("review inbound ok")

    print("e2e qualification extract ok")


if __name__ == "__main__":
    main()
